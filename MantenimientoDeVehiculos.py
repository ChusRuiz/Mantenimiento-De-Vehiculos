import openpyxl
import re
from docx import Document
from docx.shared import Inches
import mimetypes
import smtplib
import ssl
from email.message import EmailMessage
import os
from PIL import Image
from os import remove

#Verificar Agregar productos
def CoincidenciaDeDatos(agregarVehiculoSeparado):
    if ( len(agregarVehiculoSeparado) == 6 ):
        parametro = re.compile(r'\w(\w)?(\w)?(\w)?(\w)?(\w)?(\w)?(\w)?')
        coincidencia1 = parametro.search(agregarVehiculoSeparado[0])
        coincidencia2 = parametro.search(agregarVehiculoSeparado[1])
        coincidencia3 = parametro.search(agregarVehiculoSeparado[2])
        parametro = re.compile(r'(\d)?(\d)?(\d)?(\d)?\d\.\d\d')
        coincidencia4 = parametro.search(agregarVehiculoSeparado[3])
        parametro = re.compile(r'\d(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?')
        coincidencia5 = parametro.search(agregarVehiculoSeparado[4])
        coincidencia6 = parametro.search(agregarVehiculoSeparado[5])
        if (coincidencia1 == None or coincidencia2 == None or coincidencia3 == None or coincidencia4 == None or coincidencia5 == None or coincidencia6 == None):
            return False
        return True
    else:
        return False

def CoincidenciaDeCodigos(nombreImagenActual,valor):
    parametro = valor['Codigo']
    coincidencia = parametro.search(nombreImagenActual)
    if (coincidencia == None):
        return False
    return True

libro = openpyxl.load_workbook('vehiculos.xlsx')
hoja = libro['vehiculos']

#Variables
eleccion = 0
diccionarioVehiculos = {}

while eleccion != "5":
    
    vehiculos = []
    contador = 2
    contador1 = 0

    for row in range(2, hoja.max_row + 1):
        # explora fila por fila
        codigo = hoja["A" + str(row)].value
        marca = hoja["B" + str(row)].value
        modelo = hoja["C"+str(row)].value
        precio = hoja["D"+str(row)].value
        kilometraje = hoja["E"+str(row)].value
        cantidadFotografias = hoja["F"+str(row)].value
        if codigo != None and  codigo != '' :
            diccionarioVehiculos['Codigo'] = codigo
            diccionarioVehiculos['Marca'] = marca
            diccionarioVehiculos['Modelo'] = modelo
            diccionarioVehiculos['Precio'] = precio
            diccionarioVehiculos['Kilometraje'] = kilometraje
            diccionarioVehiculos['Cantidad fotografias'] = cantidadFotografias
            vehiculos.append(diccionarioVehiculos)
        diccionarioVehiculos = {}
    
    Filas = hoja.min_row + 1
    for valor in vehiculos: 
        hoja['A' + str(Filas)].value = valor['Codigo']
        hoja['B' + str(Filas)].value = valor['Marca']
        hoja['C' + str(Filas)].value = valor['Modelo']
        hoja['D' + str(Filas)].value = valor['Precio']
        hoja['E' + str(Filas)].value = valor['Kilometraje']
        hoja['F' + str(Filas)].value = valor['Cantidad fotografias']
        Filas = Filas + 1
    hoja['A' + str(Filas)].value = ''
    hoja['B' + str(Filas)].value = ''
    hoja['C' + str(Filas)].value = ''
    hoja['D' + str(Filas)].value = ''
    hoja['E' + str(Filas)].value = ''
    hoja['F' + str(Filas)].value = ''

    #guardar cambios
    libro.save("vehiculos.xlsx")

    print('Menu')
    print('1. Mantenimiento de vehículos en un archivo Excel') 
    print('2. Convertir imágenes')
    print('3. Enviar cotización') 
    print('4. Actualizar cantidad de fotografías de vehículos') 
    print('5. Salir')

    eleccion = input()

    if eleccion == "1":
        print('Menu')
        print('1. Ingresar vehículo')
        print('2. Eliminar vehículo')

        eleccion2 = input()

        if eleccion2 == "1":
            agregarVehiculo = input("Ingrese los datos en este orden Codigo,Marca,Modelo,Precio,kilometraje,cantidad de fotografias\n")
            delimitador = "," 
            agregarVehiculoSeparado = agregarVehiculo.split(delimitador)

            #verificar si existe el codigo
            for valor2 in vehiculos:
                if agregarVehiculoSeparado[0] == valor2['Codigo']:
                    print('Vehiculo ya existe')
                    contador = 1

            if CoincidenciaDeDatos(agregarVehiculoSeparado) == True and contador == 2:
                #añadir al diccionario
                diccionarioVehiculos['Codigo'] = agregarVehiculoSeparado[0]
                diccionarioVehiculos['Marca'] = agregarVehiculoSeparado[1]
                diccionarioVehiculos['Modelo'] = agregarVehiculoSeparado[2]
                diccionarioVehiculos['Precio'] = float(agregarVehiculoSeparado[3])
                diccionarioVehiculos['Kilometraje'] = float(agregarVehiculoSeparado[4])
                diccionarioVehiculos['Cantidad fotografias'] = float(agregarVehiculoSeparado[5])
                vehiculos.append(diccionarioVehiculos)
                diccionarioVehiculos = {}
  
                #añadir al libro de excel
                Filas = hoja.min_row + 1
                for valor in vehiculos: 
                    hoja['A' + str(Filas)].value = valor['Codigo']
                    hoja['B' + str(Filas)].value = valor['Marca']
                    hoja['C' + str(Filas)].value = valor['Modelo']
                    hoja['D' + str(Filas)].value = valor['Precio']
                    hoja['E' + str(Filas)].value = valor['Kilometraje']
                    hoja['F' + str(Filas)].value = valor['Cantidad fotografias']
                    Filas = Filas + 1
                #guardar cambios
                print('la informacion del vehiculo se a guardado correctamente')
                libro.save("vehiculos.xlsx")

            if CoincidenciaDeDatos(agregarVehiculoSeparado) == False:
                print("Los datos no cumplen con los parametros establecidos")
        
        #Eliminar producto
        if eleccion2 == '2':

            for valor2 in vehiculos:
                print(valor2['Codigo'])
            eliminarVehiculo = input("\nIngrese el codigo de vehiculo a eliminar\n")

            for valor2 in vehiculos:
                if eliminarVehiculo == valor2['Codigo']:
 
                    hoja['A' + str(contador)].value = None
                    hoja['B' + str(contador)].value = None
                    hoja['C' + str(contador)].value = None
                    hoja['D' + str(contador)].value = None
                    hoja['E' + str(contador)].value = None
                    hoja['F' + str(contador)].value = None
                    contador1 = 1

                    #guardar cambios
                    print('la informacion del vehiculo se a eliminado correctamente')
                    libro.save("vehiculos.xlsx")
     
                else:
                    contador = contador + 1
            if contador1 != 1:
                print('vehiculo no existe')
    
    #Convertir imagenes
    if eleccion == '2':

        nuevoAncho = int(input("Ingrese el nuevo ancho para redimensionar\n"))
        carpetaConvertir= "FOTOS"

        contenido = os.listdir(carpetaConvertir)
        imagenes = []

        for fichero in contenido:
            if os.path.isfile(os.path.join(carpetaConvertir, fichero)) and (fichero.endswith('.jpg') or fichero.endswith('.png') or fichero.endswith('.jpeg')):
                imagenes.append(fichero)

        for nombreImagenActual in imagenes:
            delimitador = "."
            nombreSeparado = nombreImagenActual.split(delimitador)
            nuevoNombre = nombreSeparado[0]+"-Convertida."+nombreSeparado[1]

            imagen = Image.open(carpetaConvertir + "\\" + nombreImagenActual)
            ancho, alto = imagen.size
            porcentajeDeDisminucion = 1 - (nuevoAncho / ancho )
            nuevoAlto = alto - (alto * porcentajeDeDisminucion)
            imagenRedimensionada = imagen.resize(( int(nuevoAncho), int(nuevoAlto)))
            imagenRedimensionada.save(carpetaConvertir + "\\"+ nuevoNombre)
            remove(carpetaConvertir + "\\"+nombreImagenActual)
        print('Se a redimencionado correctamente')

    #Cotizacion
    if eleccion == '3':
        
        for valor2 in vehiculos:
                print (valor2['Codigo'])
        codigoVehiculoCotizacion = input('\nIngrese el Codigo del vehiculo\n')
        correoCotizacion= input('\nIngrese su correo electronico\n')
        

        for valor2 in vehiculos:
                if codigoVehiculoCotizacion == valor2['Codigo']:
                    cotizacion = Document()
                    cotizacion.add_heading('Cotización', 0)
                    p = cotizacion.add_paragraph('Codigo: '+valor2['Codigo'])
                    p = cotizacion.add_paragraph('Marca: '+valor2['Marca'])
                    p = cotizacion.add_paragraph('Modelo: '+valor2['Modelo'])
                    p = cotizacion.add_paragraph('Precio: '+str(float(valor2['Precio'])))
                    p = cotizacion.add_paragraph('kilometraje: '+str(valor2['Kilometraje']))

                    carpetaUsada= "FOTOS"

                    contenido = os.listdir(carpetaUsada)
                    imagenes = []

                    for fichero in contenido:
                        if os.path.isfile(os.path.join(carpetaUsada, fichero)) and (fichero.endswith('.jpg') or fichero.endswith('.png') or fichero.endswith('.jpeg')):
                            imagenes.append(fichero)

                    for nombreImagenActual in imagenes:
                        delimitador = "-"
                        nombreSeparado = nombreImagenActual.split(delimitador)

                        if nombreSeparado[2] == codigoVehiculoCotizacion:
                            p = cotizacion.add_picture(carpetaUsada + "\\" + nombreImagenActual)

                    cotizacion.save('cotizacion.docx')
                    contador1 = 1

        if contador1 != 1:
            print('Vehiculo no existe')
        else:
            #envio por correo
            DIRECCION_DEL_SERVIDOR = "smtp.gmail.com"
            PUERTO = 587
            DIRECCION_DE_ORIGEN = "Correo de donde se va a enviar"
            CONTRASENA = 'contraseña del correo'

            #Contenido del mensaje
            mensaje = EmailMessage()
            mensaje["Subject"] = "Cotizacion"
            mensaje["From"] = DIRECCION_DE_ORIGEN
            mensaje["To"] = correoCotizacion

            mensaje.add_alternative("""
            <p> 
            <h1>No responder este mensaje</h1>
            </p>
            """, subtype = "html")

            nombre_de_archivo = "cotizacion.docx"
            ctype, encoding = mimetypes.guess_type(nombre_de_archivo)

            if ctype is None or encoding is not None:
                ctype = 'application/octet-stream'

            tipoPrincipal, subTipo = ctype.split('/', 1)

            with open(nombre_de_archivo, 'rb') as archivoLeido:
                mensaje.add_attachment(archivoLeido.read(), maintype=tipoPrincipal, subtype = subTipo, filename = nombre_de_archivo)

            context = ssl.create_default_context()

            smtp = smtplib.SMTP(DIRECCION_DEL_SERVIDOR, PUERTO)
            smtp.starttls()
            smtp.login(DIRECCION_DE_ORIGEN, CONTRASENA)
            smtp.send_message(mensaje)

            print('La cotizaciOn se a enviado correctamente\n')

    #Actualizacion de fotografias
    if eleccion == '4':
        carpetaActualizar= "FOTOS"

        contenido = os.listdir(carpetaActualizar)
        imagenes = []

        for fichero in contenido:
            if os.path.isfile(os.path.join(carpetaActualizar, fichero)) and (fichero.endswith('.jpg') or fichero.endswith('.png') or fichero.endswith('.jpeg')):
                imagenes.append(fichero)

        for  valor in vehiculos:
            contador = 0
            for nombreImagenActual in imagenes:
                delimitador = "-"
                nombreSeparado = nombreImagenActual.split(delimitador)

                if nombreSeparado[2] == valor['Codigo']:
                    contador = contador + 1

            print('Codigo: '+str(valor['Codigo']))
            print('Fotografias encontradas: '+str(contador)+'\n')

            valor['Cantidad fotografias'] = contador

            #añadir al libro de excel
        Filas = hoja.min_row + 1
        for valor in vehiculos: 
            hoja['F' + str(Filas)].value = valor['Cantidad fotografias']
            Filas = Filas + 1

        #guardar cambios
        print('la cantidad de fotografias se ha actualizado correctamente')
        libro.save("vehiculos.xlsx")
